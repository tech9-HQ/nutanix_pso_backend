from __future__ import annotations
import os
from typing import List, Optional
from docx import Document
from fastapi import UploadFile
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


async def build_combined_requirements_text(text: Optional[str], files: List[UploadFile]) -> str:
    parts: List[str] = []
    if text and str(text).strip():
        parts.append(str(text).strip())
    for f in files or []:
        try:
            blob = await f.read()
            if not blob:
                continue
            name = (f.filename or "").lower()
            if name.endswith((".txt", ".md")):
                parts.append(blob.decode("utf-8", errors="ignore"))
        except Exception:
            # ignore file read errors in short mode
            pass
    return "\n\n".join(p for p in parts if p).strip()

def safe_save_doc(document: Document, filename: str) -> str:
    outdir = os.getenv("OUTPUT_DIR", "/tmp")
    os.makedirs(outdir, exist_ok=True)
    path = os.path.join(outdir, filename)
    document.save(path)
    return path

def add_premium_cover_page(doc, company_name: str, industry: str | None, logger=None) -> None:
    """
    Renders the premium cover page used in the detailed proposal.
    Safe on failure. No return value.
    """
    try:
        accent_top = doc.add_table(rows=1, cols=1)
        accent_top.alignment = _WD_TABLE_ALIGNMENT.CENTER
        accent_cell_top = accent_top.rows[0].cells[0]
        try:
            tc = accent_cell_top._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '9500'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "2c5a7d"); tcPr.append(shd)
            trPr = accent_cell_top._element.getparent().get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), '80'); trPr.append(trHeight)
        except Exception as ex:
            if logger: logger.warning(f"Failed to style top accent: {ex}")
        accent_cell_top.text = ""
    except Exception as ex:
        if logger: logger.exception(f"Failed to create top accent: {ex}")

    from datetime import datetime
    try:
        spacer1 = doc.add_paragraph(); spacer1.paragraph_format.space_after = Pt(25)
        title_table = doc.add_table(rows=1, cols=1); title_table.alignment = _WD_TABLE_ALIGNMENT.CENTER
        title_cell = title_table.rows[0].cells[0]
        try:
            tc = title_cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '8500'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "3b76a6"); tcPr.append(shd)
            title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcMar = OxmlElement('w:tcMar')
            for k,v in {'top':'250','left':'250','bottom':'250','right':'250'}.items():
                node = OxmlElement(f'w:{k}'); node.set(qn('w:w'), v); node.set(qn('w:type'),'dxa'); tcMar.append(node)
            tcPr.append(tcMar)
        except Exception as ex:
            if logger: logger.warning(f"Failed to style title cell: {ex}")

        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.line_spacing = 1.2
        title_para.paragraph_format.space_after = Pt(0)

        r = title_para.add_run("PROFESSIONAL SERVICES\n"); r.bold = True; r.font.size = Pt(32); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(255,255,255)
        r = title_para.add_run("Statement of Work\n"); r.font.size = Pt(20); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(240,248,255)
        r = title_para.add_run(f"{datetime.now().year}"); r.font.size = Pt(12); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(200,220,240)
    except Exception as ex:
        if logger: logger.exception(f"Failed to create title section: {ex}")

    try:
        spacer2 = doc.add_paragraph(); spacer2.paragraph_format.space_after = Pt(30)
        client_table = doc.add_table(rows=1, cols=1); client_table.alignment = _WD_TABLE_ALIGNMENT.CENTER
        client_cell = client_table.rows[0].cells[0]
        try:
            tc = client_cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '7500'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "f8fbfd"); tcPr.append(shd)
            tcBorders = OxmlElement('w:tcBorders')
            top_border = OxmlElement('w:top'); top_border.set(qn('w:val'),'single'); top_border.set(qn('w:sz'),'20'); top_border.set(qn('w:color'),'3b76a6'); tcBorders.append(top_border)
            for border in ['left','bottom','right']:
                b = OxmlElement(f'w:{border}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:color'),'cccccc'); tcBorders.append(b)
            tcPr.append(tcBorders)
            tcMar = OxmlElement('w:tcMar')
            for k,v in {'top':'200','left':'200','bottom':'200','right':'200'}.items():
                node = OxmlElement(f'w:{k}'); node.set(qn('w:w'), v); node.set(qn('w:type'),'dxa'); tcMar.append(node)
            tcPr.append(tcMar)
        except Exception as ex:
            if logger: logger.warning(f"Failed to style client cell: {ex}")

        p = client_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("PREPARED FOR\n"); r.font.size = Pt(8); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(120,120,120); r.font.all_caps = True
        r = p.add_run("•\n"); r.font.size = Pt(9); r.font.color.rgb = RGBColor(59,118,166)
        r = p.add_run(f"{company_name}\n"); r.bold = True; r.font.size = Pt(24); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(59,118,166)
        if industry:
            r = p.add_run("—\n"); r.font.size = Pt(10); r.font.color.rgb = RGBColor(150,150,150)
            r = p.add_run(f"{industry}"); r.font.size = Pt(13); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(100,100,100)
    except Exception as ex:
        if logger: logger.exception(f"Failed to create client section: {ex}")

    try:
        spacer3 = doc.add_paragraph(); spacer3.paragraph_format.space_after = Pt(20)
        meta_table = doc.add_table(rows=3, cols=2); meta_table.alignment = _WD_TABLE_ALIGNMENT.CENTER
        from datetime import datetime as _dt
        current_date = _dt.now().strftime("%B %d, %Y")
        try:
            tbl = meta_table._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            if tbl.tblPr is None: tbl.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for b in ['top','left','bottom','right','insideH','insideV']:
                bn = OxmlElement(f'w:{b}'); bn.set(qn('w:val'),'none'); tblBorders.append(bn)
            tblPr.append(tblBorders)
        except Exception as ex:
            if logger: logger.warning(f"Failed to hide meta table borders: {ex}")
        meta_table.rows[0].cells[0].text = "Document Date:"; meta_table.rows[0].cells[1].text = current_date
        meta_table.rows[1].cells[0].text = "Prepared By:";   meta_table.rows[1].cells[1].text = "Integrated Tech9 Labs Pvt. Ltd."
        meta_table.rows[2].cells[0].text = "Version:";       meta_table.rows[2].cells[1].text = "1.0"
        for row in meta_table.rows:
            try:
                trPr = row._element.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'),'200'); trPr.append(trHeight)
            except: 
                pass
            lc = row.cells[0]; vc = row.cells[1]
            for paragraph in lc.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = RGBColor(100,100,100); run.font.italic = True
            for paragraph in vc.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = RGBColor(59,118,166); run.bold = True
    except Exception as ex:
        if logger: logger.exception(f"Failed to create metadata section: {ex}")

    try:
        spacer4 = doc.add_paragraph(); spacer4.paragraph_format.space_before = Pt(20)
        tagline = doc.add_paragraph(); tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER; tagline.paragraph_format.space_after = Pt(15)
        r = tagline.add_run("Delivering Excellence in Professional Services"); r.italic = True; r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(120,120,120)
        for i, color in enumerate(['d6e9f5','3b76a6','2c5a7d']):
            accent_table = doc.add_table(rows=1, cols=1); accent_table.alignment = _WD_TABLE_ALIGNMENT.CENTER
            accent_cell = accent_table.rows[0].cells[0]
            try:
                tc = accent_cell._tc; tcPr = tc.get_or_add_tcPr()
                widths = ['9000','7500','6000']
                tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), widths[i]); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), color); tcPr.append(shd)
                heights = ['30','45','60']; trPr = accent_cell._element.getparent().get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), heights[i]); trPr.append(trHeight)
            except Exception as ex:
                if logger: logger.warning(f"Failed to style accent bar {i}: {ex}")
            accent_cell.text = ""
    except Exception as ex:
        if logger: logger.exception(f"Failed to create footer section: {ex}")

    doc.add_page_break()