# app/routers/generate_proposal.py
from __future__ import annotations

import os, re, json, asyncio, logging
from decimal import Decimal
from typing import List, Dict, Any, Optional

import requests
from fastapi import APIRouter, Request, Form, File, UploadFile, HTTPException
from fastapi.responses import FileResponse

# router first (good)
router = APIRouter(prefix="", tags=["proposal"])

# docx libs (heavy, CPU-bound, synchronous)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# project deps
# NOTE: imports below expose module-level state (PDF_CHUNKS, fitz) and network clients — be cautious of side effects
from app.services.proposals_repo import (
    supabase_ro, load_pdf_chunks, REFERENCE_PDF_PATH, PDF_CHUNKS, fitz
)
from app.services.llm_selector import (
    _llm_generate_introduction,
    _llm_generate_executive_summary,
    _llm_generate_solution_summary,
    _llm_generate_milestone_plan,
    _llm_generate_acceptance_criteria,
    generate_service_content,
    create_comprehensive_fallback_content,
    generate_services_parallel,
    generate_project_governance,
)
from app.services.journey import map_kb_and_pdf_chunks_for_service
from app.utils.text import sanitize_filename, sanitize_bm_name
from app.utils.timeit import add_bookmark_to_paragraph
from app.config import TERMS_AND_CONDITIONS
from app.models.schemas import ProposalRequest

# logger setup — OK but prefer app-wide logging config
logger = logging.getLogger("generate_proposal")
if not logger.handlers:
    _h = logging.StreamHandler()
    _h.setFormatter(logging.Formatter("[%(levelname)s] %(name)s: %(message)s"))
    logger.addHandler(_h)
logger.setLevel(logging.INFO)


# -------------------------
# small helpers
# -------------------------
def _clean_duplicate_headings(text: str, heading_to_remove: str) -> str:
    """
    Remove repeated headings and divider lines from LLM output.
    Good defensive cleanup before adding to doc.
    """
    if not text:
        return text
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.lower() == heading_to_remove.lower():
            continue
        if re.match(r'^-{3,}$', stripped):
            continue
        lines.append(line)
    return '\n'.join(lines)


def _fallback_exec_summary(company: str, req: Optional[str], deployment: Optional[str]) -> str:
    """
    Local deterministic fallback when LLM executive summary fails.
    Keep this in sync with templates used elsewhere.
    """
    req = (req or "Requirements will be finalized during discovery.").strip()
    dep = (deployment or "N/A").strip()
    return (
        "Executive Summary\n"
        f"This Statement of Work describes professional services for {company}.\n\n"
        "Client Overview\n"
        f"{company} seeks predictable delivery and measurable outcomes via a structured engagement.\n\n"
        "Requirement Summary\n"
        f"{req}\n\n"
        "Assumptions\n"
        "- Stakeholders are available for decisions\n"
        "- Timely access to required environments and subscriptions\n"
        "- Work proceeds in agreed sprints\n"
        f"- Deployment model: {dep}\n"
    )


def _distribute_days(total_days: int) -> list[tuple[str,int,str]]:
    """
    Deterministic split of total_days across phases.
    - Guarantees sum(raw) == total_days
    - Ensures nonzero phases get at least 1 day when total_days > 0
    Remarks:
    - Business rule is embedded here. If this changes, move weights to config.
    """
    total_days = max(int(total_days or 0), 0)
    if total_days == 0:
        return [
            ("Discovery & Planning", 1, "Scope confirmation, access, success criteria"),
            ("Build & Configure", 1, "Implement baseline, integrations"),
            ("Validate & Handover", 1, "UAT, documentation, sign-off"),
        ]
    phases = [
        ("Discovery & Planning", 0.25, "Scope confirmation, access, success criteria"),
        ("Build & Configure",   0.45, "Implement baseline, integrations"),
        ("Testing & UAT",       0.20, "Functional, security, and performance validation"),
        ("Go-Live & Handover",  0.10, "Cutover, KT, final sign-off"),
    ]
    raw = [max(0, round(total_days * w)) for _, w, _ in phases]
    diff = total_days - sum(raw)
    idxs = sorted(range(len(raw)), key=lambda i: -phases[i][1])
    i = 0
    while diff != 0 and idxs:
        j = idxs[i % len(idxs)]
        if diff > 0:
            raw[j] += 1; diff -= 1
        else:
            if raw[j] > 0:
                raw[j] -= 1; diff += 1
        i += 1
    for k in range(len(raw)):
        if total_days > 0 and raw[k] == 0:
            raw[k] = 1
    while sum(raw) > total_days:
        k = min((ii for ii,v in enumerate(raw) if v > 1), key=lambda ii: raw[ii], default=None)
        if k is None: break
        raw[k] -= 1
    while sum(raw) < total_days:
        k = max(range(len(raw)), key=lambda ii: raw[ii])
        raw[k] += 1
    return [(phases[i][0], raw[i], phases[i][2]) for i in range(len(phases))]


@router.post("/generate_proposal")
async def generate_proposal(
    request: Request,
    company_name: Optional[str] = Form(None),
    client_requirements: Optional[str] = Form(None),
    industry: Optional[str] = Form(None),
    deployment_type: Optional[str] = Form(None),
    runtime_pdf: Optional[str] = Form("false"),
    services: Optional[str] = Form("[]"),
    files: List[UploadFile] = File([]),
):
    """
    Generates a DOCX SOW (detailed).
    IMPORTANT: This handler is async but performs many synchronous, blocking operations:
      - python-docx manipulation (CPU-bound)
      - persistent DB calls via supabase_ro (blocking)
      - LLM calls invoked via asyncio.to_thread (ok) but may still be heavy
      - requests.get for FX rate (blocking)
    Production guidance: offload to worker or run blocking parts with run_in_threadpool()
    """
    try:
        # accept JSON body if form-data empty
        is_services_empty = (not services) or (isinstance(services, str) and services.strip() in ("", "[]"))
        has_input = any([
            bool(company_name and str(company_name).strip()),
            bool(client_requirements and str(client_requirements).strip()),
            bool(industry and str(industry).strip()),
            bool(deployment_type and str(deployment_type).strip()),
            bool(files and len(files) > 0),
            not is_services_empty
        ])

        if not has_input:
            # fallback JSON body ingestion — good for clients that send JSON
            try:
                payload = await request.json()
                company_name = payload.get("company_name") or company_name
                client_requirements = payload.get("client_requirements") or client_requirements
                industry = payload.get("industry") or industry
                deployment_type = payload.get("deployment_type") or deployment_type
                runtime_pdf = payload.get("runtime_pdf", runtime_pdf)
                services = json.dumps(payload.get("services", []))
            except Exception:
                # explicit 400 for missing content — avoid returning vague errors later
                raise HTTPException(status_code=400, detail="Empty request: provide company_name, requirements, files or services")

        # normalize inputs
        company_name = (company_name or "Client").strip()
        runtime_pdf_bool = str(runtime_pdf).lower() in ("1", "true", "yes", "on")
        try:
            parsed_services = json.loads(services or "[]")
            if isinstance(parsed_services, dict):
                parsed_services = [parsed_services]
            if not isinstance(parsed_services, list):
                parsed_services = []
        except Exception:
            parsed_services = []

        # validate/shape into Pydantic ProposalRequest early — good practice
        req_dict = {
            "company_name": company_name,
            "client_requirements": client_requirements,
            "industry": industry,
            "deployment_type": deployment_type,
            "runtime_pdf": runtime_pdf_bool,
            "services": parsed_services,
        }
        try:
            req = ProposalRequest(**req_dict)
        except Exception as e:
            # Validate early and return 400 rather than failing deep in flow
            raise HTTPException(status_code=400, detail=f"Invalid request fields: {e}")

        # combine requirements + text files
        # This reads uploaded files into memory — validate file sizes to avoid OOM
        async def _build_combined_requirements_text(text: Optional[str], files: List[UploadFile]) -> str:
            parts: List[str] = []
            if text and str(text).strip():
                parts.append(str(text).strip())
            for f in files or []:
                try:
                    content = await f.read()
                    if not content:
                        continue
                    name = (f.filename or "").lower()
                    if name.endswith((".txt", ".md")):
                        parts.append(content.decode("utf-8", errors="ignore"))
                except Exception:
                    logger.exception("Failed reading uploaded file")
            return "\n\n".join(p for p in parts if p).strip()

        combined_requirements = await _build_combined_requirements_text(getattr(req, "client_requirements", None), files)
        if not combined_requirements:
            combined_requirements = " ".join([p for p in [getattr(req, "industry", None) or "", getattr(req, "deployment_type", None) or ""] if p]).strip()

        if not combined_requirements and not (getattr(req, "services", None) or []):
            raise HTTPException(status_code=400, detail="No requirements text, files, or services provided.")

        # runtime PDF reload — be careful: mutating global PDF_CHUNKS is racy in multi-worker deployments
        global PDF_CHUNKS
        if runtime_pdf_bool and REFERENCE_PDF_PATH and (fitz is not None) and os.path.exists(REFERENCE_PDF_PATH):
            try:
                # load_pdf_chunks is likely blocking and potentially slow — consider run_in_threadpool or worker
                PDF_CHUNKS = load_pdf_chunks(REFERENCE_PDF_PATH)
            except Exception:
                logger.exception("Runtime PDF reload attempted but failed")

        # parallel static-section generation (LLM calls) using asyncio.to_thread — good approach for sync LLM functions
        INTRO_TIMEOUT = int(os.getenv("INTRO_TIMEOUT", "45"))
        EXEC_TIMEOUT = int(os.getenv("EXEC_TIMEOUT", "120"))  # increased
        SOL_TIMEOUT = int(os.getenv("SOL_TIMEOUT", "75"))
        SERVICES_TIMEOUT = int(os.getenv("SERVICES_TIMEOUT", "300"))

        # spawn blocking LLM ops in threads to avoid blocking the event loop
        intro_task = asyncio.to_thread(
            _llm_generate_introduction,
            req.company_name,
            req.industry,
            combined_requirements,
            req.deployment_type
        )
        exec_task = asyncio.to_thread(
            _llm_generate_executive_summary,
            req.company_name,
            combined_requirements,
            req.deployment_type
        )
        sol_task = asyncio.to_thread(
            _llm_generate_solution_summary,
            req.company_name,
            combined_requirements,
            req.services,
            req.industry,
            req.deployment_type
        )
        services_task = asyncio.to_thread(generate_services_parallel, parsed_services or [], req)

        intro_text = exec_text = sol_text = None
        service_contents: Dict[str, str] = {}

        # Await with timeouts — good. But fallback behaviors and logging are important (you have them).
        try:
            intro_text = await asyncio.wait_for(intro_task, timeout=INTRO_TIMEOUT)
        except asyncio.TimeoutError:
            logger.warning("Intro generation timed out")
        except Exception:
            logger.exception("Intro generation failed")

        try:
            exec_text = await asyncio.wait_for(exec_task, timeout=EXEC_TIMEOUT)
        except asyncio.TimeoutError:
            logger.warning("Executive summary generation timed out")
        except Exception:
            logger.exception("Executive summary generation failed")

        # enforce fallback if missing — good defensive pattern
        if not exec_text or not exec_text.strip():
            exec_text = _fallback_exec_summary(req.company_name, combined_requirements, req.deployment_type)

        try:
            sol_text = await asyncio.wait_for(sol_task, timeout=SOL_TIMEOUT)
        except asyncio.TimeoutError:
            logger.warning("Solution summary generation timed out")
        except Exception:
            logger.exception("Solution summary generation failed")

        try:
            svc_map = await asyncio.wait_for(services_task, timeout=SERVICES_TIMEOUT)
            if isinstance(svc_map, dict):
                service_contents = svc_map
            else:
                service_contents = {}
        except asyncio.TimeoutError:
            logger.warning("Service generation timed out")
            service_contents = {}
        except Exception:
            logger.exception("Service generation failed")
            service_contents = {}

        # build doc (python-docx synchronous, heavy)
        doc = Document()
        bookmark_map: Dict[str, str] = {}
        used_bookmarks = set()

        def get_unique_bookmark(base_name: str, prefix: str = "bm") -> str:
            bm = sanitize_bm_name(base_name, prefix); counter = 1; original_bm = bm
            while bm in used_bookmarks:
                bm = f"{original_bm}_{counter}"; counter += 1
            used_bookmarks.add(bm); return bm

        # cover page
        # NOTE: all these docx style manipulations are synchronous and may be slow for many concurrent requests
        try:
            accent_top = doc.add_table(rows=1, cols=1)
            accent_top.alignment = _WD_TABLE_ALIGNMENT.CENTER
            accent_cell_top = accent_top.rows[0].cells[0]
            try:
                tc = accent_cell_top._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '9500'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "2c5a7d"); tcPr.append(shd)
                trPr = accent_cell_top._element.getparent().get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight'); trHeight.set(qn('w:val'), '80'); trPr.append(trHeight)
            except Exception as ex:
                # styling errors are non-fatal — log and continue
                logger.warning(f"Failed to style top accent: {ex}")
            accent_cell_top.text = ""
        except Exception as ex:
            logger.exception(f"Failed to create top accent: {ex}")

        from datetime import datetime
        try:
            spacer1 = doc.add_paragraph(); spacer1.paragraph_format.space_after = Pt(25)
            title_table = doc.add_table(rows=1, cols=1); title_table.alignment = _WD_TABLE_ALIGNMENT.CENTER
            title_cell = title_table.rows[0].cells[0]
            try:
                tc = title_cell._tc; tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '8500'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "3b76a6"); tcPr.append(shd)
                title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tcMar = OxmlElement('w:tcMar')
                for k,v in {'top':'250','left':'250','bottom':'250','right':'250'}.items():
                    node = OxmlElement(f'w:{k}'); node.set(qn('w:w'), v); node.set(qn('w:type'),'dxa'); tcMar.append(node)
                tcPr.append(tcMar)
            except Exception as ex:
                logger.warning(f"Failed to style title cell: {ex}")

            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.line_spacing = 1.2
            title_para.paragraph_format.space_after = Pt(0)

            r = title_para.add_run("PROFESSIONAL SERVICES\n"); r.bold = True; r.font.size = Pt(32); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(255,255,255)
            r = title_para.add_run("Statement of Work\n"); r.font.size = Pt(20); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(240,248,255)
            r = title_para.add_run(f"{datetime.now().year}"); r.font.size = Pt(12); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(200,220,240)
        except Exception as ex:
            logger.exception(f"Failed to create title section: {ex}")

        # [snip — many docx layout manipulations follow; they are mostly styling - continue same pattern]
        # ... (the rest of your doc composition logic remains unchanged; keep the same defensive try/except blocks)

        # Table of contents generation uses bookmarks and internal hyperlinks.
        # NOTE: Word bookmarks are fragile; ensure bookmark names are unique and short.

        # Introduction
        intro_heading = doc.add_heading("Introduction", level=1)
        add_bookmark_to_paragraph(intro_heading, bookmark_map.get("Introduction", "introduction"), {"id": 1})
        if intro_text:
            intro_text = _clean_duplicate_headings(intro_text, "Introduction")
            cleaned_lines = []
            for line in intro_text.splitlines():
                l = line.strip()
                if not l:
                    continue
                if re.fullmatch(r"[*_]*\s*introduction\s*[*_]*:?\s*", l, flags=re.IGNORECASE):
                    continue
                cleaned_lines.append(l)
            for para in "\n".join(cleaned_lines).split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            doc.add_paragraph("Introduction will be finalized during project initiation.")
        doc.add_page_break()

        # Executive Summary — cleaned and styled
        exec_heading = doc.add_heading("Executive Summary", level=1)
        add_bookmark_to_paragraph(exec_heading, bookmark_map.get("Executive Summary", "executive_summary"), {"id": 2})
        if exec_text:
            exec_text = _clean_duplicate_headings(exec_text, "Executive Summary")
            for line in exec_text.splitlines():
                ln = line.strip()
                if not ln:
                    continue
                if re.fullmatch(r"\s*executive\s*summary\s*:?\s*", ln, flags=re.IGNORECASE):
                    continue
                if re.match(r'^(Client Overview|Requirement Summary|Detailed Requirements|Assumptions)\s*$', ln, re.IGNORECASE):
                    p = doc.add_paragraph(); r = p.add_run(ln); r.bold = True; r.font.color.rgb = RGBColor(59,118,166); r.font.size = Pt(13)
                    continue
                if ln.startswith("- "):
                    doc.add_paragraph(ln[2:], style="List Bullet"); continue
                doc.add_paragraph(ln)
        else:
            doc.add_paragraph(f"This proposal outlines the Professional Services engagement for {req.company_name}.")
        doc.add_page_break()

        # Solution Summary and Mapping — parse LLM table lines if present, otherwise write free text
        sol_heading = doc.add_heading("Solution Summary and Mapping", level=1)
        add_bookmark_to_paragraph(sol_heading, bookmark_map.get("Solution Summary and Mapping", "solution_summary"), {"id": 3})
        if sol_text:
            sol_text = _clean_duplicate_headings(sol_text, "Solution Summary and Mapping")

            def _strip_md(s: str) -> str:
                s = s.strip()
                s = re.sub(r"^[*_`]+|[*_`]+$", "", s)
                return re.sub(r"\s+", " ", s).strip()

            header_aliases = {
                "client requirement", "requirement", "requirements",
                "proposed solution / approach", "solution", "proposed solution", "approach",
                "expected outcome", "outcome", "result"
            }

            rows, seen_children = [], set()

            def split_composite(name: str):
                n = name or ""
                for sep in (" and ", " & ", " / ", " + "):
                    if sep in n.lower():
                        return [p.strip() for p in re.split(r"\s+(?:and|&|/|\+)\s+", n, flags=re.I) if p.strip()]
                return []

            for l in [x.strip() for x in sol_text.splitlines() if x.strip()]:
                if re.fullmatch(r"[-_]{3,}", l):
                    continue
                if "|" in l and l.count("|") >= 2:
                    parts = [ _strip_md(p) for p in l.strip().strip("|").split("|") ]
                    lowers = [ p.lower() for p in parts ]
                    if all(p in header_aliases for p in lowers if p):
                        continue
                    if all(re.fullmatch(r"[:\-\._]{3,}", p) for p in parts):
                        continue
                    while len(parts) < 3:
                        parts.append("")
                    req_cell, sol_cell, out_cell = parts[0], parts[1], parts[2]
                    children = split_composite(sol_cell)
                    if children:
                        seen_children.update(c.lower() for c in children)
                    if sol_cell.strip().lower() in seen_children and not children:
                        continue
                    rows.append((req_cell, sol_cell, out_cell))
                    continue
                if l.lower() in {"solution summary and mapping", "overview", "requirement-to-solution mapping"}:
                    continue
                doc.add_paragraph(l)

            if rows:
                tbl = doc.add_table(rows=1, cols=3); tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells; hdr[0].text = "Client Requirement"; hdr[1].text = "Proposed Solution / Approach"; hdr[2].text = "Expected Outcome"
                for r in rows:
                    c = tbl.add_row().cells; c[0].text, c[1].text, c[2].text = r[0], r[1], r[2]
        doc.add_page_break()

        # Milestone Plan — deterministic distribution across phases
        milestone_heading = doc.add_heading("Milestone Plan", level=1)
        add_bookmark_to_paragraph(milestone_heading, bookmark_map.get("Milestone Plan", "milestone_plan"), {"id": 4})

        total_days = sum([int(getattr(s, "duration", 0) or (s.get("duration", 0) if isinstance(s, dict) else 0)) for s in (getattr(req, "services", []) or [])])
        duration_summary = f"Approx. {total_days} days total" if total_days else "TBD"

        # deterministic table from selected services
        rows = _distribute_days(total_days or 0)
        milestone_table = doc.add_table(rows=1, cols=3); milestone_table.style = "Table Grid"
        hdr = milestone_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Phase", "Key Activities", "Duration (days)"
        for name, days, acts in rows:
            c = milestone_table.add_row().cells
            c[0].text, c[1].text, c[2].text = name, acts, str(days)

        p = doc.add_paragraph(); p.add_run(f"Overall timeline: {duration_summary}").italic = True
        doc.add_page_break()

        # Service sections + pricing capture
        summary_rows: List[Dict[str, Any]] = []
        from collections import OrderedDict
        services_by_category = OrderedDict(); ordered_categories: List[str] = []

        # group services by category preserving order
        for item in parsed_services:
            if isinstance(item, dict):
                svc_name = (item.get("service_name") or "Service").strip()
                cat_name = (item.get("category") or "Uncategorized").strip()
            else:
                svc_name = (getattr(item, "service_name", None) or "Service").strip()
                cat_name = (getattr(item, "category", None) or "Uncategorized").strip()
            if cat_name not in services_by_category:
                services_by_category[cat_name] = []; ordered_categories.append(cat_name)
            services_by_category[cat_name].append(item)

        for cidx, cat_name in enumerate(ordered_categories):
            cat_heading = doc.add_heading(cat_name, level=1)
            add_bookmark_to_paragraph(cat_heading, bookmark_map.get(cat_name, sanitize_bm_name(cat_name, "bm_cat")), {"id": 5})

            for item in services_by_category.get(cat_name, []):
                if isinstance(item, dict):
                    svc_name = (item.get("service_name") or "Service").strip()
                    duration_val = int(item.get("duration") or 1)
                    service_comment = item.get("comment") or ""
                else:
                    svc_name = (getattr(item, "service_name", None) or "Service").strip()
                    duration_val = int(getattr(item, "duration", None) or 1)
                    service_comment = getattr(item, "comment", None) or ""

                svc_heading = doc.add_heading(svc_name, level=2)
                svc_key = f"{cat_name}|{svc_name}"
                add_bookmark_to_paragraph(svc_heading, bookmark_map.get(svc_key, sanitize_bm_name(svc_key, "bm_svc")), {"id": 6})

                # DB lookup (blocking) — consider caching or moving this into threadpool
                service_info = None
                try:
                    resp = supabase_ro.table("proposals").select("*").eq("service_name", svc_name).limit(1).execute()
                    if resp and getattr(resp, "data", None):
                        service_info = resp.data[0]
                except Exception:
                    logger.exception("DB lookup failed for service %s", svc_name)

                # KB + PDF pointers (likely cheap) then content generation
                kb_pointers = map_kb_and_pdf_chunks_for_service(svc_name, top_k=8)

                gen = service_contents.get(svc_name)
                if not gen:
                    try:
                        gen = generate_service_content(
                            svc_name, service_info, kb_pointers, duration_val,
                            industry=getattr(req, "industry", None),
                            service_comment=service_comment,
                            deployment_type=getattr(req, "deployment_type", None),
                        )
                    except Exception:
                        logger.exception("Service content generation failed for %s", svc_name)
                        gen = create_comprehensive_fallback_content(
                            svc_name, (service_info.get("positioning") if service_info else ""), duration_val,
                        )

                if gen:
                    gen = _clean_duplicate_headings(gen, svc_name)
                    seen_lines = set(); current_section = None
                    for raw_line in gen.splitlines():
                        line = raw_line.strip()
                        if not line or line in seen_lines:
                            continue
                        seen_lines.add(line)
                        if line.rstrip(":") in {"Executive Summary","Objective","Scope of Work","Deliverables","Assumptions & Dependencies"}:
                            current_section = line.rstrip(":"); doc.add_heading(current_section, level=3); continue
                        if line.startswith("- "):
                            doc.add_paragraph(line[2:], style="List Bullet"); continue
                        if line.startswith("Duration:"):
                            p = doc.add_paragraph(); p.add_run(line).bold = True; continue
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph(f"{svc_name} – content unavailable.")

                # price parsing — use Decimal for money; handle malformed values defensively
                price_per_day = Decimal("0.00")
                try:
                    if service_info:
                        pval = service_info.get("price_man_day") or service_info.get("price")
                        if pval is not None:
                            price_per_day = Decimal(str(pval))
                except Exception:
                    logger.exception("Failed to parse price for %s", svc_name)

                total_cost = (price_per_day * Decimal(duration_val)).quantize(Decimal("0.01"))

                summary_rows.append({
                    "category": cat_name,
                    "service": svc_name,
                    "duration": duration_val,
                    "cost": f"{price_per_day:.2f}",
                    "total_cost": f"{total_cost:.2f}",
                })

            if cidx < len(ordered_categories) - 1:
                doc.add_page_break()

        doc.add_page_break()

        # Acceptance Criteria — generate with LLM call with fallback
        ac_heading = doc.add_heading("Acceptance Criteria", level=1)
        add_bookmark_to_paragraph(ac_heading, bookmark_map.get("Acceptance Criteria", "acceptance_criteria"), {"id": 7})
        try:
            ac_text = _llm_generate_acceptance_criteria(req.company_name, req.services, req.deployment_type)
            if ac_text:
                ac_text = _clean_duplicate_headings(ac_text, "Acceptance Criteria")
                criteria, seen = [], set()
                for l in ac_text.splitlines():
                    t = l.strip()
                    if not t:
                        continue
                    if re.fullmatch(r"[-_]{3,}", t):
                        continue
                    if t.lower() in {"overview", "criteria", "acceptance criteria"}:
                        continue
                    if t.startswith("- "):
                        t = t[2:].strip()
                    if t and t not in seen:
                        seen.add(t); criteria.append(t)

                if criteria:
                    table = doc.add_table(rows=1, cols=2); table.style = "Light Grid Accent 1"
                    try:
                        tbl = table._element; tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                        if tbl.tblPr is None: tbl.insert(0, tblPr)
                        tblBorders = OxmlElement('w:tblBorders')
                        for border in ['top','left','right']:
                            bn = OxmlElement(f'w:{border}'); bn.set(qn('w:val'), 'none'); tblBorders.append(bn)
                        insideH = OxmlElement('w:insideH'); insideH.set(qn('w:val'),'single'); insideH.set(qn('w:sz'),'6'); insideH.set(qn('w:color'),'CCCCCC'); tblBorders.append(insideH)
                        insideV = OxmlElement('w:insideV'); insideV.set(qn('w:val'),'none'); tblBorders.append(insideV)
                        bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:color'),'CCCCCC'); tblBorders.append(bottom)
                        tblPr.append(tblBorders)
                    except Exception as ex:
                        logger.warning(f"Failed to customize table borders: {ex}")

                    hdr = table.rows[0].cells; hdr[0].text = "Criteria"; hdr[1].text = "Yes / No"
                    # styling omitted here for brevity — your implementation applies styling
                    for crit in criteria:
                        row = table.add_row().cells
                        row[0].text = crit; row[1].text = "☐"
                        # style rows
                    table.columns[0].width = Inches(5.5); table.columns[1].width = Inches(1.0)
        except Exception:
            logger.exception("Acceptance criteria generation error")
            doc.add_paragraph("Acceptance criteria will be finalized during discovery.")
        doc.add_page_break()

        # Project Governance — generate and sanitize LLM output
        gov_heading = doc.add_heading("Project Governance", level=1)
        add_bookmark_to_paragraph(gov_heading, bookmark_map.get("Project Governance", "project_governance"), {"id": 8})
        try:
            svc_names = ", ".join([str(getattr(s, "service_name", "") or (s.get("service_name","") if isinstance(s, dict) else "")) for s in (req.services or [])][:8])
            gov_text = generate_project_governance(req.company_name, req.industry, services_overview=svc_names, deployment_type=req.deployment_type)
            client = (req.company_name or "Client").strip(); vendor = "Integrated Tech9 Labs Pvt. Ltd."

            def _natural_join(items):
                parts = [x.strip() for x in items if x and x.strip()]
                if not parts: return ""
                if len(parts) == 1: return parts[0]
                return ", ".join(parts[:-1]) + " and " + parts[-1]

            svc_list = [x.strip() for x in (svc_names.split(",") if svc_names else []) if x.strip()]
            lead = (
                f"At {vendor}, we prioritize robust project governance to ensure seamless execution and "
                f"alignment with {client}’s objectives. Our structured approach fosters transparency, "
                f"accountability, and timely decision-making throughout the project lifecycle."
            )
            lead += f" We are committed to delivering {_natural_join(svc_list)}." if svc_list else " We are committed to delivering the agreed scope with measurable outcomes."

            # sanitize placeholders in gov_text — good defensive step
            for pat in (r"\[\s*Your Company Name\s*\]", r"\(\s*Your Company Name\s*\)", r"\bYour Company Name\b", r"\byour\s+company\b"):
                gov_text = re.sub(pat, vendor, gov_text, flags=re.IGNORECASE)
            gov_text = re.sub(r"\b(the\s+customer|customer)\b", client, gov_text, flags=re.IGNORECASE)
            gov_text = re.sub(r"^\s*At\s+[^,\n]+,\s+we[^\n]*\n?", "", gov_text, flags=re.IGNORECASE)
            gov_text = _clean_duplicate_headings(gov_text, "Project Governance")
            gov_text = "\n".join(ln for ln in (gov_text.splitlines() if gov_text else []) if not re.fullmatch(r"[-_]{3,}", ln.strip())).strip()

            doc.add_paragraph(lead)
            for line in gov_text.splitlines():
                ls = line.strip()
                if not ls:
                    continue
                if ls.lower().startswith(("project governance", "overview")):
                    continue
                if ls.startswith("- "):
                    doc.add_paragraph(ls[2:], style="List Bullet")
                else:
                    doc.add_paragraph(ls)
        except Exception:
            logger.exception("Project governance generation failed")
            doc.add_paragraph(f"Project governance will be agreed with {req.company_name} during project initiation.")
        doc.add_page_break()

        # Cost Summary
        cost_heading = doc.add_heading("Cost Summary", level=1)
        try:
            add_bookmark_to_paragraph(cost_heading, bookmark_map.get("Cost Summary", "cost_summary"), {"id": 9})
        except Exception:
            pass

        intro_para = doc.add_paragraph()
        intro_para.add_run("The following cost breakdown reflects the professional services outlined in this Statement of Work. All pricing is provided in both USD and INR for transparency and clarity.")
        intro_para.paragraph_format.space_after = Pt(12)

        # FX rate fetch (blocking) — cache this value for several minutes/hrs to avoid network calls per request
        usd_inr_rate = None
        try:
            r = requests.get("https://api.exchangerate.host/latest", params={"base": "USD", "symbols": "INR"}, timeout=5)
            jr = r.json()
            if jr and jr.get("rates") and jr["rates"].get("INR"):
                usd_inr_rate = Decimal(str(jr["rates"]["INR"]))
        except Exception:
            logger.exception("FX fetch failed; using fallback rate")
        if not usd_inr_rate:
            usd_inr_rate = Decimal(os.getenv("FALLBACK_USD_INR", "87.95"))

        cost_table = doc.add_table(rows=1, cols=6); cost_table.style = "Table Grid"
        headers = ["Category","Service","Duration\n(days)","Rate Per Day\n(USD)","Subtotal\n(USD)","Subtotal\n(INR)"]
        for i, h in enumerate(headers):
            hdr_cell = cost_table.rows[0].cells[i]; hdr_cell.text = h; hdr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in hdr_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(255,255,255)
            try:
                shading_elm = parse_xml(r'<w:shd {} w:fill="3b76a6"/>'.format(nsdecls('w')))
                hdr_cell._element.get_or_add_tcPr().append(shading_elm)
            except Exception as ex:
                logger.warning(f"Failed to add header background for cell {i}: {ex}")
        try:
            cost_table.columns[0].width = Inches(1.3); cost_table.columns[1].width = Inches(2.2)
            cost_table.columns[2].width = Inches(0.9); cost_table.columns[3].width = Inches(1.0)
            cost_table.columns[4].width = Inches(1.0); cost_table.columns[5].width = Inches(1.1)
        except Exception as ex:
            logger.warning(f"Failed to set column widths: {ex}")

        grand_total_usd = Decimal("0.00"); grand_total_inr = Decimal("0.00")
        for idx, rdata in enumerate(summary_rows):
            try:
                row_cells = cost_table.add_row().cells
                row_cells[0].text = str(rdata.get("category","") or ""); row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT; row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[1].text = str(rdata.get("service","") or ""); row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT; row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row_cells[2].text = str(rdata.get("duration","") or ""); row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                usd_cost = Decimal(str(rdata.get("cost","0") or "0")); row_cells[3].text = f"${usd_cost:,.2f}"; row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                usd_total = Decimal(str(rdata.get("total_cost","0") or "0")); row_cells[4].text = f"${usd_total:,.2f}"; row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                try:
                    inr_total = (usd_total * usd_inr_rate).quantize(Decimal("0.01"))
                except Exception:
                    inr_total = (usd_total * Decimal("87.95")).quantize(Decimal("0.01"))
                row_cells[5].text = f"₹{inr_total:,.2f}"; row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for cell in row_cells:
                    for p in cell.paragraphs:
                        for r in p.runs: r.font.size = Pt(10); r.font.name = "Calibri"
                if idx % 2 == 1:
                    try:
                        for cell in row_cells:
                            shading_elm = parse_xml(r'<w:shd {} w:fill="f8fbfd"/>'.format(nsdecls('w')))
                            cell._element.get_or_add_tcPr().append(shading_elm)
                    except: pass
                grand_total_usd += usd_total; grand_total_inr += inr_total
            except Exception:
                logger.exception("Failed writing a cost row; skipping row")

        # Subtotal
        try:
            subtotal_row = cost_table.add_row().cells
            subtotal_row[0].text = ""; subtotal_row[1].text = "Subtotal"; subtotal_row[2].text = ""; subtotal_row[3].text = ""
            subtotal_row[4].text = f"${grand_total_usd:,.2f}"; subtotal_row[5].text = f"₹{grand_total_inr:,.2f}"
            for i, cell in enumerate(subtotal_row):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    for r in p.runs: r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"
            try:
                for cell in subtotal_row:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="d6e9f5"/>'.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(shading_elm)
            except: pass
        except Exception:
            logger.exception("Failed to append subtotal row")

        # GST
        try:
            gst_rate_pct = Decimal(os.getenv("GST_RATE_PCT", "18"))
            gst_amount_inr = (grand_total_inr * gst_rate_pct / Decimal("100")).quantize(Decimal("0.01"))
            gst_row = cost_table.add_row().cells
            gst_row[0].text = ""; gst_row[1].text = f"GST @ {gst_rate_pct}%"; gst_row[2].text = ""; gst_row[3].text = ""
            gst_row[4].text = "—"; gst_row[5].text = f"₹{gst_amount_inr:,.2f}"
            for i, cell in enumerate(gst_row):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size = Pt(10); r.font.name = "Calibri"; r.italic = True
        except Exception:
            logger.exception("Failed to append GST row")

        # Grand Total
        try:
            final_inr = (grand_total_inr + gst_amount_inr).quantize(Decimal("0.01"))
            total_row = cost_table.add_row().cells
            total_row[0].text = ""; total_row[1].text = "Grand Total (Incl. GST)"; total_row[2].text = ""; total_row[3].text = ""
            total_row[4].text = f"${grand_total_usd:,.2f}"; total_row[5].text = f"₹{final_inr:,.2f}"
            for i, cell in enumerate(total_row):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(59,118,166)
            try:
                for cell in total_row:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="e8f4f8"/>'.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(shading_elm)
            except: pass
        except Exception:
            logger.exception("Failed to append grand total row")

        # INR amount in words — local helper for Indian numbering; prefer tested library for i18n
        def _indian_words(n: int) -> str:
            if n == 0: return "Zero"
            ones = ["","One","Two","Three","Four","Five","Six","Seven","Eight","Nine"]
            tens = ["","Ten","Twenty","Thirty","Forty","Fifty","Sixty","Seventy","Eighty","Ninety"]
            teens = ["Ten","Eleven","Twelve","Thirteen","Fourteen","Fifteen","Sixteen","Seventeen","Eighteen","Nineteen"]
            def two_digits(x):
                if x < 10: return ones[x]
                if x < 20: return teens[x-10]
                return tens[x//10] + (" " + ones[x%10] if x%10 else "")
            def three_digits(x):
                h, r = x // 100, x % 100
                if h and r: return ones[h] + " Hundred " + two_digits(r)
                if h: return ones[h] + " Hundred"
                return two_digits(r)
            parts = []; crore = n // 10000000; n %= 10000000
            lakh = n // 100000; n %= 100000
            thousand = n // 1000; n %= 1000
            hundred = n
            if crore: parts.append(three_digits(crore) + " Crore")
            if lakh: parts.append(three_digits(lakh) + " Lakh")
            if thousand: parts.append(three_digits(thousand) + " Thousand")
            if hundred: parts.append(three_digits(hundred))
            return " ".join(parts)

        doc.add_paragraph("")
        try:
            inr_whole = int(final_inr.quantize(Decimal("1"), rounding="ROUND_DOWN"))
            inr_frac = int(((final_inr - Decimal(inr_whole)) * 100).quantize(Decimal("1")))
            rupees_words = _indian_words(inr_whole)
            amount_para = doc.add_paragraph(); amount_para.paragraph_format.space_before = Pt(6)
            r = amount_para.add_run("Total Payable (Incl. GST): "); r.bold = True; r.font.size = Pt(10); r.font.name = "Calibri"
            amount_words = f"INR {inr_whole:,} ({rupees_words} Rupees"
            if inr_frac: amount_words += f" and {inr_frac:02d} Paise"
            amount_words += ") only"
            r = amount_para.add_run(amount_words); r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(59,118,166)
        except Exception:
            logger.exception("Failed to write INR total in words")

        doc.add_paragraph("")
        note_para = doc.add_paragraph(); note_para.paragraph_format.space_before = Pt(12)
        r = note_para.add_run("Note: "); r.bold = True; r.font.size = Pt(9); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(100,100,100)
        r = note_para.add_run("All line rates are exclusive of GST. The Grand Total includes applicable GST as computed above. Exchange rate applied: "); r.font.size = Pt(9); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(100,100,100)
        r = note_para.add_run(f"1 USD = ₹{usd_inr_rate}"); r.font.size = Pt(9); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(100,100,100); r.italic = True

        doc.add_page_break()

        # Terms & Conditions
        terms_heading = doc.add_heading("Terms and Conditions", level=1)
        add_bookmark_to_paragraph(terms_heading, bookmark_map.get("Terms and Conditions", "terms_and_conditions"), {"id": 10})
        for t in TERMS_AND_CONDITIONS:
            doc.add_paragraph(t, style="List Bullet")

        # save and return (writes to disk)
        def _safe_save_doc(document: Document, filename: str) -> str:
            outdir = os.getenv("OUTPUT_DIR", "/tmp")
            os.makedirs(outdir, exist_ok=True)
            path = os.path.join(outdir, filename)
            document.save(path)
            return path

        safe_name = sanitize_filename(req.company_name or "document")
        filename = f"{safe_name}_Proposal.docx"
        filepath = _safe_save_doc(doc, filename)
        logger.info("Proposal generated at %s", filepath)

        return FileResponse(
            path=filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(filepath),
        )

    except HTTPException:
        raise
    except Exception as e:
        # Avoid returning full exception text to end users in production (PII/leaks)
        logger.exception("Error in generate_proposal endpoint: %s", str(e))
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")
