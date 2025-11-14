# app/routers/generate_proposal_short.py

from __future__ import annotations
import os, re, json, datetime
from decimal import Decimal
from typing import List, Optional, Dict, Any
import requests
import logging

from fastapi import APIRouter, Request, Form, File, UploadFile, HTTPException
from fastapi.responses import FileResponse

# python-docx imports (heavy, CPU-bound, blocks event loop)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Modular helpers (good separation)
from app.utils.doc_helpers import (
    build_combined_requirements_text,
    safe_save_doc,
    add_premium_cover_page,
)
from app.utils.text import sanitize_filename
from app.services.journey import map_kb_and_pdf_chunks_for_service
from app.services.llm_selector import (
    generate_service_content,
    create_comprehensive_fallback_content,
    _llm_generate_introduction,
)
from app.config import TERMS_AND_CONDITIONS

# Use proposals_repo as a module so we can update shared globals safely
# NOTE: this mutates module-level globals (PDF_CHUNKS etc). That can be racy under multiple workers.
import app.services.proposals_repo as repo

logger = logging.getLogger("generate_proposal_short")
router = APIRouter(prefix="", tags=["proposal"])


# -----------------------------
# Local helpers
# -----------------------------
def _get_price_per_day(svc_info: Dict[str, Any] | None, payload: Dict[str, Any] | None) -> Decimal:
    """
    Normalize price-per-day from various candidate fields. Returns Decimal.
    - Accepts values as numbers or strings (with commas).
    - Returns Decimal('0.00') on failure.
    Notes:
    - Using Decimal is correct for money, but ensure a consistent Decimal context
      (rounding mode) across the app when doing math.
    """
    cands: List[Any] = []
    if isinstance(svc_info, dict):
        cands += [
            svc_info.get("price_man_day"),
            svc_info.get("price"),
            svc_info.get("price_per_day"),
            svc_info.get("cost_per_day"),
        ]
    if isinstance(payload, dict):
        cands += [
            payload.get("price_man_day"),
            payload.get("price"),
            payload.get("price_per_day"),
            payload.get("cost_per_day"),
        ]
    for c in cands:
        if c is None:
            continue
        try:
            val = Decimal(str(c))
            if val >= 0:
                return val
        except Exception:
            continue
    return Decimal("0.00")


def _number_to_words(num: int) -> str:
    """
    Convert integer to words (English). Handles up to trillions.
    - Keep this utility simple; for localization consider `num2words` library.
    """
    units = ["","One","Two","Three","Four","Five","Six","Seven","Eight","Nine"]
    teens = ["Ten","Eleven","Twelve","Thirteen","Fourteen","Fifteen","Sixteen","Seventeen","Eighteen","Nineteen"]
    tens  = ["","","Twenty","Thirty","Forty","Fifty","Sixty","Seventy","Eighty","Ninety"]
    thousands = ["","Thousand","Million","Billion","Trillion"]
    if num == 0:
        return "Zero"

    def chunk_words(n: int) -> str:
        out: List[str] = []
        if n >= 100:
            out += [units[n // 100], "Hundred"]
            n %= 100
        if n >= 20:
            out.append(tens[n // 10])
            if n % 10:
                out.append(units[n % 10])
        elif n >= 10:
            out.append(teens[n - 10])
        elif n > 0:
            out.append(units[n])
        return " ".join(w for w in out if w)

    parts: List[str] = []
    i = 0
    while num:
        n = num % 1000
        if n:
            cw = chunk_words(n)
            parts.insert(0, f"{cw} {thousands[i]}".strip())
        num //= 1000
        i += 1
    return ", ".join(parts)


def _derive_benefit(text: str | None) -> str:
    """
    Heuristic extraction of a 'benefit' sentence from text.
    - Rewrites first matching sentence to address the customer ("you").
    - If nothing found, uses first paragraph clause or a default.
    Notes:
    - Heuristics like this are fine for a short proposal but can be brittle;
      consider using an LLM summarization or a small rule-based extractor with tests.
    """
    if not text:
        return "You will benefit from reduced risk and faster delivery."
    for s in re.split(r'(?<=[.!?])\s+', text.replace("\r", "\n")):
        low = s.lower()
        if any(k in low for k in ["reduce","minimiz","improv","accelerat","ensure","enable","mitigate","increase","speed","de-risk"]):
            s = re.sub(r'^\s*[-•\*]\s*', '', s).strip().rstrip('.')
            s = re.sub(r'\bwe\b', 'you', s, flags=re.I)
            s = re.sub(r'\bour\b', 'your', s, flags=re.I)
            return f"You will benefit from {s}."
    first = re.split(r'\n\s*\n|\.\s+', text.strip())[0].strip().rstrip('.')
    if not first:
        return "You will benefit from reduced risk and faster delivery."
    return f"You will benefit from {first[:180]}. This improves delivery predictability and reduces risk."


def _sanitize(txt: str | None) -> str:
    """
    Clean up generated text:
    - collapse repeated lines, remove repeated headings, strip extraneous whitespace
    - remove common section headings that might appear in LLM output
    """
    if not txt:
        return ""
    out: List[str] = []
    prev: Optional[str] = None
    for l in txt.splitlines():
        s = l.rstrip()
        if s and s == prev:
            continue
        out.append(s)
        prev = s
    s = "\n".join(out)
    s = re.sub(r'\n{3,}', '\n\n', s)
    s = re.sub(r'(?m)^\s*(Executive Summary|Introduction|Overview)\s*[:\-–]?\s*$', '', s)
    return s.strip()


# -----------------------------
# Endpoint
# -----------------------------
@router.post("/generate_proposal_short")
async def generate_proposal_short(
    request: Request,
    company_name: Optional[str] = Form(None),
    client_requirements: Optional[str] = Form(None),
    industry: Optional[str] = Form(None),
    deployment_type: Optional[str] = Form(None),
    runtime_pdf: Optional[str] = Form("false"),
    services: Optional[str] = Form("[]"),
    files: List[UploadFile] = File([]),
):
    """
    Concise one-pager-per-service proposal. Premium cover page. Robust fallbacks. Returns DOCX.

    Important operational notes:
    - This handler uses heavy synchronous libraries (python-docx, requests, repo.load_pdf_chunks).
      It's declared `async` but performs blocking work — you should run blocking parts in a threadpool
      (e.g., `starlette.concurrency.run_in_threadpool`) or move generation to a background worker.
    - File sizes and number of files should be limited (validate UploadFile.size / headers). Otherwise
      memory / disk exhaustion is possible.
    - LLM/HTTP calls (generate_service_content, exchange-rate API) must have timeouts and circuit-breakers.
    """
    try:
        # Accept JSON body too
        services_empty = not services or (isinstance(services, str) and services.strip() in ("", "[]"))
        if not any([company_name, client_requirements, industry, deployment_type, files, not services_empty]):
            try:
                body = await request.json()
                company_name = body.get("company_name")
                client_requirements = body.get("client_requirements")
                industry = body.get("industry")
                deployment_type = body.get("deployment_type")
                runtime_pdf = body.get("runtime_pdf", "false")
                services = json.dumps(body.get("services", []))
            except Exception:
                # Use explicit HTTPException for a clear 400 response
                raise HTTPException(status_code=400, detail="Empty or invalid request body")

        company_name = (company_name or "Client").strip()
        runtime_pdf_bool = str(runtime_pdf).lower() in ("1", "true", "yes", "on")

        # Parse services - validate shape early to avoid failing mid-generation
        try:
            parsed_services = json.loads(services or "[]")
            if isinstance(parsed_services, dict):
                parsed_services = [parsed_services]
            if not isinstance(parsed_services, list):
                parsed_services = []
        except Exception:
            parsed_services = []

        # Build combined requirements (text + files)
        # NOTE: build_combined_requirements_text likely reads files into memory — validate size.
        combined_req = await build_combined_requirements_text(client_requirements, files)
        if not combined_req:
            combined_req = " ".join([p for p in [industry or "", deployment_type or ""] if p]).strip()
        if not combined_req and not parsed_services:
            raise HTTPException(status_code=400, detail="No requirements text, files, or services provided.")

        # Optional runtime PDF reload (update shared cache in repo)
        # WARNING: Mutating module-level globals in `repo` can be racy when using multiple worker processes.
        # If you need runtime reloads in production, use a synchronization mechanism or a central store (Redis/DB).
        if runtime_pdf_bool and repo.REFERENCE_PDF_PATH and repo.fitz and os.path.exists(repo.REFERENCE_PDF_PATH):
            try:
                logger.info("Runtime PDF reload (short)")
                # This can be slow / blocking — consider run_in_threadpool or background task
                repo.PDF_CHUNKS = repo.load_pdf_chunks(repo.REFERENCE_PDF_PATH)
            except Exception:
                logger.exception("Runtime PDF reload failed (short)")

        # Create document (python-docx is sync & CPU-bound)
        # Consider generating documents in a worker (Celery/RQ) for better scalability.
        doc = Document()

        # Premium cover identical to detailed proposal
        # add_premium_cover_page likely manipulates doc object and may add images — ensure images are sized and safe.
        add_premium_cover_page(doc, company_name, industry, logger)  # ends with page break

        # Context + date
        # Use UTC-consistent date or localize as needed for clients.
        doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d-%b-%Y')}")
        ctx = f"Based on {company_name}'s {(deployment_type or 'deployment')} objectives"
        if industry:
            ctx += f" in the {industry} sector"
        if combined_req:
            head = combined_req.strip().splitlines()[0][:120]
            if head:
                ctx += f', and the key requirement: "{head.strip()}"'
        ctx += ", the following concise service recommendations and costs are proposed."
        doc.add_paragraph(ctx)

        # Introduction (short) — LLM-style generation; ensure timeouts and sanitized inputs
        try:
            intro = _llm_generate_introduction(company_name, industry, combined_req, deployment_type)
        except Exception:
            intro = None
        if intro:
            seg = " ".join(re.split(r'(?<=[.!?])\s+', intro.strip())[:2])
            seg = re.sub(r'\bwe\b', 'you', seg, flags=re.I)
            doc.add_paragraph(seg[:1000])
        else:
            doc.add_paragraph(f"This short proposal outlines recommended services for {company_name}.")

        # Services loop — this is the main CPU + I/O hotspot
        summary_rows: List[Dict[str, Any]] = []
        grand_total = Decimal("0.00")

        for item in parsed_services:
            # tolerate both dicts and objects — but prefer a validated schema earlier
            if isinstance(item, dict):
                svc_name = item.get("service_name") or item.get("service") or "Service"
                category = item.get("category") or "Uncategorized"
                days = int(item.get("duration") or item.get("default_days") or 1)
                comment = item.get("comment") or ""
                payload = item
            else:
                svc_name = getattr(item, "service_name", None) or getattr(item, "service", None) or "Service"
                category = getattr(item, "category", None) or "Uncategorized"
                days = int(getattr(item, "duration", None) or getattr(item, "default_days", None) or 1)
                comment = getattr(item, "comment", None) or ""
                payload = {}

            doc.add_heading(svc_name, level=2)

            # DB lookup — synchronous call to repo.supabase_ro. If you use networked DB, this blocks.
            svc_info: Dict[str, Any] | None = None
            try:
                resp = repo.supabase_ro.table("proposals").select("*").eq("service_name", svc_name).limit(1).execute()
                svc_info = (getattr(resp, "data", None) or [None])[0]
            except Exception:
                # Log DB errors but continue — fallback content will be used.
                logger.exception("DB query failed for %s", svc_name)

            # Content generation — may call LLMs, lookup KB, etc. Must have precautions:
            # - timeouts, retries
            # - input sanitization
            # - cost/latency controls
            try:
                content = generate_service_content(
                    svc_name,
                    svc_info,
                    map_kb_and_pdf_chunks_for_service(svc_name, top_k=4),
                    days,
                    industry=industry,
                    service_comment=comment,
                    deployment_type=deployment_type,
                )
            except Exception:
                # Robust fallbacks: positioning from DB or synthesized fallback content
                if isinstance(svc_info, dict) and (svc_info.get("positioning") or svc_info.get("description")):
                    content = svc_info.get("positioning") or svc_info.get("description")
                else:
                    content = create_comprehensive_fallback_content(
                        svc_name,
                        (svc_info.get("positioning") if isinstance(svc_info, dict) else ""),
                        days,
                    )

            content = _sanitize(content) or f"{svc_name} — professional services engagement. Estimated duration: {days} day(s)."

            # Render top nuggets (limit to first 4 lines to control doc length)
            doc.add_paragraph("Overview:")
            shown = 0
            for ln in [l.strip() for l in content.splitlines() if l.strip()]:
                if shown >= 4:
                    break
                if ln.endswith(":"):
                    doc.add_heading(ln[:-1], level=3)
                elif ln.startswith("- "):
                    doc.add_paragraph(ln[2:], style="List Bullet")
                else:
                    doc.add_paragraph(ln)
                shown += 1

            doc.add_paragraph(f"Estimated Duration: {days} day{'s' if days > 1 else ''}")
            doc.add_paragraph(f"Key benefit: {_derive_benefit(content)}")

            # Pricing: use Decimal arithmetic; ensure consistent currency
            price = _get_price_per_day(svc_info if isinstance(svc_info, dict) else None, payload)
            total = (price * Decimal(days)).quantize(Decimal("0.01"))
            summary_rows.append({
                "category": category,
                "service": svc_name,
                "duration": days,
                "cost": f"{price:.2f}",
                "total_cost": f"{total:.2f}",
            })
            grand_total += total

        # Cost summary table (python-docx table operations are fine here but slow for many rows)
        doc.add_heading("Cost Summary", level=1)

        # FX fetch — using requests (sync). Use async httpx or single threaded offload.
        usd_inr = None
        try:
            r = requests.get("https://api.exchangerate.host/latest", params={"base": "USD", "symbols": "INR"}, timeout=5)
            j = r.json()
            if j and j.get("rates") and j["rates"].get("INR"):
                usd_inr = Decimal(str(j["rates"]["INR"]))
        except Exception:
            logger.exception("FX fetch failed (short)")
        if not usd_inr:
            # Fallback value from env — ensure it's present and sensible
            usd_inr = Decimal(os.getenv("FALLBACK_USD_INR", "87.95"))

        # Build a table and populate rows
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        hdrs = ["Category", "Service", "Duration (days)", "Cost (per day) (USD)", "Total cost (USD)", "Total cost (INR)"]
        for i, h in enumerate(hdrs):
            c = tbl.rows[0].cells[i]
            c.text = h
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True

        grand_inr = Decimal("0.00")
        for row in summary_rows:
            rc = tbl.add_row().cells
            rc[0].text = row["category"]
            rc[1].text = row["service"]
            rc[2].text = str(row["duration"])
            usd_cost = Decimal(row["cost"])
            rc[3].text = f"${usd_cost:,.2f}"
            usd_total = Decimal(row["total_cost"])
            rc[4].text = f"${usd_total:,.2f}"
            inr_total = (usd_total * usd_inr).quantize(Decimal("0.01"))
            rc[5].text = f"₹{inr_total:,.2f}"
            grand_inr += inr_total

        tr = tbl.add_row().cells
        tr[0].text = ""
        tr[1].text = "GRAND TOTAL"
        tr[2].text = ""
        tr[3].text = ""
        tr[4].text = f"${grand_total:,.2f}"
        tr[5].text = f"₹{grand_inr:,.2f}"
        for c in tr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True

        # Amount in words — careful with large numbers and localization
        try:
            whole = int(grand_inr.quantize(Decimal("1"), rounding="ROUND_DOWN"))
            frac = int(((grand_inr - Decimal(whole)) * 100).quantize(Decimal("1")))
            words = _number_to_words(whole)
            s = f"INR {whole:,} ({words} Rupees"
            if frac:
                s += f" and {frac:02d} Paise"
            s += ") only."
            p = doc.add_paragraph()
            p.add_run("Grand Total (INR) in words: ").bold = True
            p.add_run(s)
        except Exception:
            # Ignore word conversion issues — not critical
            pass

        # Terms
        doc.add_heading("Terms & Conditions", level=1)
        for t in TERMS_AND_CONDITIONS:
            doc.add_paragraph(t, style="List Bullet")

        # Closing
        doc.add_paragraph("We appreciate your consideration and look forward to the opportunity to partner with you.")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("Integrated Tech9 Labs Pvt. Ltd.")

        # Save + return
        # Use sanitized filename and safe_save_doc ensures unique path, but validate path and permissions.
        filename = f"{sanitize_filename(company_name)}_Short_Proposal.docx"
        path = safe_save_doc(doc, filename)
        logger.info("Short proposal generated: %s", path)

        # Returning a file from disk: consider scheduling deletion of the file after send (BackgroundTasks)
        # or streaming bytes directly to avoid storing on disk for long.
        return FileResponse(
            path=path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(path),
        )

    except HTTPException:
        # Preserve explicit HTTPExceptions
        raise
    except Exception as e:
        # Avoid returning str(e) in production (can leak internals). Log full stack and give sanitized message.
        logger.exception("Error in generate_proposal_short")
        raise HTTPException(status_code=500, detail=f"Server error")
